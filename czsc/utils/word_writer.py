# -*- coding: utf-8 -*-
"""
author: zengbin93
email: zeng_bin8888@163.com
create_dt: 2021/10/28 21:22
describe: 实现一个用python-docx写word文档的辅助工具

参考资料：
1. https://cloud.tencent.com/developer/article/1512325
2. https://blog.csdn.net/zhouz92/article/details/107066709
"""
import os
import docx
import pandas as pd
try:
    from docx.oxml.ns import qn
    from docx.shared import Cm, Pt, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
except ImportError:
    raise ImportError("请安装 python-docx 库")


class WordWriter:
    """用 Word 文档记录信息"""

    def __init__(self, file_docx=None):
        self.file_docx = file_docx
        if file_docx and os.path.exists(file_docx):
            self.document = docx.Document(file_docx)
        else:
            self.document = docx.Document()
        self.document.core_properties.author = "Reporter"

        # 设置正文样式
        self.document.styles["Normal"].font.name = 'Times New Roman'
        self.document.styles["Normal"].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def add_title(self, text):
        self.document.core_properties.title = text
        title_ = self.document.add_heading(level=0)
        title_.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        title_run = title_.add_run(text)
        title_run.font.size = Pt(22)
        title_run.font.bold = True
        title_run.font.name = 'Times New Roman'
        title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    def add_heading(self, text, level=1):
        if level == 1:
            size = 18
        elif level == 2:
            size = 16
        else:
            size = 14
        title_ = self.document.add_heading(level=level)
        title_run = title_.add_run(text)
        title_run.font.size = Pt(size)
        title_run.font.name = 'Times New Roman'
        title_run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        title_run.font.color.rgb = RGBColor(0, 0, 0)

    def add_paragraph(self, text, style=None, bold=False, first_line_indent=0.74):
        """新增段落

        :param text: 文本
        :param style: 段落样式
        :param bold: 是否加粗
        :param first_line_indent: 首行缩进，0.74 表示两个空格
        :return:
        """
        p = self.document.add_paragraph(style=style)
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(8)

        text = p.add_run(text)
        text.bold = bold
        text.font.name = 'Times New Roman'
        text.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        text.font.size = Pt(12)

    def add_df_table(self, df: pd.DataFrame, style='Table Grid', **kwargs):
        """添加数据表

        https://www.jianshu.com/p/93e0df92cf16

        :param df: 数据表
        :param style: 表格样式
        :return:
        """
        if df.empty:
            print("add_df_table error: 传入的数据表是空的")
            return

        table = self.document.add_table(rows=1, cols=df.shape[1], style=style)
        # 设置整个表格字体属性
        table.style.font.size = Pt(kwargs.get("font_size", 10))
        table.style.font.color.rgb = RGBColor(0, 0, 0)
        table.style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        records = df.to_dict("records")
        columns = df.columns.to_list()

        hdr_cells = table.rows[0].cells
        for i, c in enumerate(columns):
            hdr_cells[i].text = c

        for row in records:
            row_cells = table.add_row().cells
            for i, c in enumerate(columns):
                row_cells[i].text = str(row[c])

    def add_picture(self, file, width=None, height=None, alignment='center') -> None:
        """写入图片到文档中

        :param file: 图片文件路径
        :param width: 图片宽度，默认单位 cm
        :param height: 图片高度，默认单位 cm
        :param alignment: 图片对齐，默认 center
        :return:
        """
        alignment_map = {
            'center': WD_PARAGRAPH_ALIGNMENT.CENTER,
            'left': WD_PARAGRAPH_ALIGNMENT.LEFT,
            'right': WD_PARAGRAPH_ALIGNMENT.RIGHT,
        }
        if isinstance(width, int):
            width = Cm(width)

        if isinstance(height, int):
            height = Cm(height)

        paragraph = self.document.add_paragraph()
        paragraph.alignment = alignment_map[alignment]
        run = paragraph.add_run("")
        run.add_picture(file, width, height)

    def add_page_break(self):
        """添加分页符"""
        self.document.add_page_break()

    def save(self, file_docx=None):
        """保存结果到文件"""
        if not file_docx:
            file_docx = self.file_docx
        self.document.save(file_docx)


def test_word_writer():
    import os
    import pandas as pd
    import inspect
    import matplotlib.pyplot as plt
    
    reporter = WordWriter()

    reporter.add_title('Reporter测试记录文档')
    reporter.add_paragraph('这个方法可以用来生成完整的回测报告。以文字、图表为主，统一格式')

    reporter.add_heading('一、研究背景介绍', level=1)
    reporter.add_paragraph('python-docx 是用于创建和更新Microsoft Word（.docx）文件的Python库。')

    reporter.add_heading('1) 无序项目', level=2)
    reporter.add_paragraph('无序项目1', style='List Bullet')
    reporter.add_paragraph('无序项目2', style='List Bullet')
    reporter.add_paragraph('无序项目3', style='List Bullet')

    reporter.add_heading('2) 有序项目', level=2)
    reporter.add_paragraph('有序项目1', style='List Number')
    reporter.add_paragraph('有序项目2', style='List Number')
    reporter.add_paragraph('有序项目3', style='List Number')
    reporter.add_page_break()

    reporter.add_heading('二、主要研究结果', level=1)
    reporter.add_paragraph('Python中可以用docx来生成word文档，docx中可以自定义文字的大小和字体等。')
    reporter.add_paragraph("段落是Word中的一个块级对象，在其所在容器的左右边界内显示文本，当文本超过"
                           "右边界时自动换行。段落的边界通常是页边界，也可以是分栏排版时的栏边界，或者"
                           "表格单元格中的边界。段落格式用于控制段落在其容器（例如页、栏、单元格）中的"
                           "布局，例如对齐方式、左缩进、右缩进、首行缩进、行距、段前距离、段后距离、换"
                           "页方式、Tab键字符格式等。")

    reporter.add_paragraph("""
    newfile = docx.Document()
    newfile.styles['Normal'].font.name = 'Times New Roman'
    newfile.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    """, style='Normal')

    reporter.add_heading('1) 无序项目', level=2)
    reporter.add_paragraph('无序项目1', style='List Bullet')
    reporter.add_paragraph('无序项目2', style='List Bullet')
    reporter.add_paragraph('无序项目3', style='List Bullet')

    reporter.add_heading('2) 有序项目', level=2)
    reporter.add_paragraph('有序项目1', style='List Number')
    reporter.add_paragraph('有序项目2', style='List Number')
    reporter.add_paragraph('有序项目3', style='List Number')

    # 测试表格写入
    df = pd.DataFrame([{'x': 1, 'y': 2, 'z': 3}, {'x': 1, 'y': 2, 'z': 3}, {'x': 1, 'y': 2, 'z': 3}])
    reporter.add_df_table(df)
    reporter.add_paragraph('\n\n')
    reporter.add_df_table(df, 'Light Grid Accent 2')
    reporter.add_paragraph('\n\n')
    reporter.add_df_table(df, 'Colorful Grid Accent 2')
    reporter.add_paragraph('\n\n')
    reporter.add_df_table(df, 'Colorful Shading Accent 2')
    reporter.add_paragraph('\n\n')
    reporter.add_df_table(df, 'Dark List Accent 2')
    reporter.add_paragraph('\n\n')
    reporter.add_df_table(df, 'Medium List 1 Accent 2')
    reporter.save("reporter_test.docx")

    # 测试续写文档
    file_docx = "reporter_test.docx"
    reporter = WordWriter(file_docx)
    reporter.add_page_break()
    reporter.add_heading('三、讨论', level=1)
    reporter.add_paragraph('Python中可以用docx来生成word文档，docx中可以自定义文字的大小和字体等。')
    reporter.add_paragraph("""段落是Word中的一个块级对象，在其所在容器的左右边界内显示文本，当文本超过
    右边界时自动换行。段落的边界通常是页边界，也可以是分栏排版时的栏边界，或者表格单元格中的边界。段落格式
    用于控制段落在其容器（例如页、栏、单元格）中的布局，例如对齐方式、左缩进、右缩进、首行缩进、行距、段前
    距离、段后距离、换页方式、Tab键字符格式等。
        """.strip().replace("\n", ""))

    reporter.add_page_break()
    reporter.add_heading('四、源码', level=1)
    reporter.add_paragraph(inspect.getsource(WordWriter))

    reporter.save(file_docx)

    # 写入图片
    plt.close()
    df = pd.DataFrame({'x': list(range(100)), 'y': list(range(100))})
    ax1 = plt.subplot(211)
    ax1.plot(df['x'], df['y'], 'go--', linewidth=2, markersize=12)
    ax2 = plt.subplot(212)
    ax2.plot(df['x'], df['y'], 'ro--', linewidth=2, markersize=12)
    file_png = "x.png"
    plt.savefig(file_png)
    reporter.add_picture(file_png)
    reporter.add_picture(file_png, width=8, height=6)
    reporter.save(file_docx)
    os.remove(file_png)

    # 查看全部样式
    all_styles = list(reporter.document.styles.__iter__())
    print(all_styles)
    os.remove(file_docx)
